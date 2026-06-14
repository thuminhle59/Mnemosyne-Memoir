"""Render a MeetingReport to .docx (python-docx) or .pdf (weasyprint via HTML)."""
import io
import re
from html import escape
from docx import Document
from models import MeetingReport


def build_filename(title: str, date: str, fmt: str) -> str:
    """Sanitize a report file name (shared by the HTTP response and email attachment)."""
    safe = re.sub(r"[^\w\- ]", "", title).strip().replace(" ", "_") or "meeting"
    return f"{safe}_{date}.{fmt}"


def filename(r: MeetingReport, fmt: str) -> str:
    return build_filename(r.title, r.date, fmt)


def render_docx(r: MeetingReport) -> bytes:
    doc = Document()
    doc.add_heading(r.title, level=0)
    doc.add_paragraph(f"Ngày: {r.date}" + (f" · Thời lượng: ~{r.duration_min} phút" if r.duration_min else ""))

    doc.add_heading("Tóm tắt", level=1)
    doc.add_paragraph(r.summary)

    if r.key_points:
        doc.add_heading("Ý chính", level=1)
        for k in r.key_points:
            doc.add_paragraph(k, style="List Bullet")

    if r.decisions:
        doc.add_heading("Quyết định", level=1)
        for d in r.decisions:
            doc.add_paragraph(d.text, style="List Number")

    if r.action_items:
        doc.add_heading("Action items", level=1)
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = "Việc", "Người phụ trách", "Deadline", "Ưu tiên"
        for ai in r.action_items:
            row = table.add_row().cells
            row[0].text = ai.task
            row[1].text = ai.owner or "-"
            row[2].text = ai.deadline or "-"
            row[3].text = ai.priority

    if r.risks:
        doc.add_heading("Rủi ro / Blocker", level=1)
        for x in r.risks:
            doc.add_paragraph(x, style="List Bullet")

    if r.open_questions:
        doc.add_heading("Vấn đề còn treo", level=1)
        for q in r.open_questions:
            doc.add_paragraph(q, style="List Bullet")

    if r.next_meeting:
        doc.add_heading("Buổi họp tiếp theo", level=1)
        doc.add_paragraph(r.next_meeting)

    doc.add_heading("Transcript đầy đủ", level=1)
    doc.add_paragraph(r.full_transcript)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _html(r: MeetingReport) -> str:
    def ul(items):
        return "<ul>" + "".join(f"<li>{escape(i)}</li>" for i in items) + "</ul>" if items else ""
    dur = f" · Thời lượng: ~{r.duration_min} phút" if r.duration_min else ""

    key_points_html = f"<h2>Ý chính</h2>{ul(r.key_points)}" if r.key_points else ""
    decisions_html = (
        "<h2>Quyết định</h2><ol>"
        + "".join(f"<li>{escape(d.text)}</li>" for d in r.decisions)
        + "</ol>"
    ) if r.decisions else ""

    rows = "".join(
        f"<tr><td>{escape(ai.task)}</td><td>{escape(ai.owner or '-')}</td>"
        f"<td>{escape(ai.deadline or '-')}</td><td>{escape(ai.priority)}</td></tr>"
        for ai in r.action_items
    )
    action_html = (
        "<h2>Action items</h2>"
        "<table border='1' cellspacing='0' cellpadding='4'>"
        "<tr><th>Việc</th><th>Người phụ trách</th><th>Deadline</th><th>Ưu tiên</th></tr>"
        f"{rows}</table>"
    ) if r.action_items else ""

    risks_html = f"<h2>Rủi ro / Blocker</h2>{ul(r.risks)}" if r.risks else ""
    open_q_html = f"<h2>Vấn đề còn treo</h2>{ul(r.open_questions)}" if r.open_questions else ""
    next_html = f"<h2>Buổi họp tiếp theo</h2><p>{escape(r.next_meeting)}</p>" if r.next_meeting else ""

    return f"""<!doctype html><html><head><meta charset="utf-8">
<style>body{{font-family:'DejaVu Sans',sans-serif;font-size:12px}} h1{{font-size:20px}}
td,th{{font-size:11px}}</style></head><body>
<h1>{escape(r.title)}</h1><p>Ngày: {escape(r.date)}{dur}</p>
<h2>Tóm tắt</h2><p>{escape(r.summary)}</p>
{key_points_html}
{decisions_html}
{action_html}
{risks_html}
{open_q_html}
{next_html}
<h2>Transcript đầy đủ</h2><pre style="white-space:pre-wrap">{escape(r.full_transcript)}</pre>
</body></html>"""


def render_pdf(r: MeetingReport) -> bytes:
    from weasyprint import HTML  # imported lazily so docx path works without weasyprint libs
    return HTML(string=_html(r)).write_pdf()
